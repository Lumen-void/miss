from __future__ import annotations

import json
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image


ROOT = Path("/Applications/XAMPP/xamppfiles/htdocs/miss")
DATA = json.loads((ROOT / "storage/run-1-architecture-data.json").read_text())
SHOT_DIR = ROOT / "storage/architecture-screenshots-doc"
OUT = ROOT / "docs/MIS_Tool_Architecture_Run_1.docx"


SCREENSHOTS = [
    ("01-start-dashboard.png", "Start - Command Center", "The tool starts at the command center. This page shows the active monthly run, key MIS totals, health/readiness, and the entry points into import, sales, inventory, MIS review, and reports."),
    ("02-dashboard-close-board.png", "Dashboard - Close Board", "This explains close readiness and what must be reviewed before finalization."),
    ("03-dashboard-trends.png", "Dashboard - Trends", "This compares the current run against prior runs where data exists."),
    ("04-import-guided-sources.png", "Import - Guided Sources", "This is the starting point for data capture. It supports browser-assisted portal import, API import, manual upload, and scheduled runs."),
    ("05-import-api-connect.png", "Import - API Connect", "This connects API sources such as website/custom API, Shopify, WooCommerce, Amazon SP-API, and Zoho P&L mapping."),
    ("06-import-activity.png", "Import - Activity", "This shows the job log for browser automation and whether imports completed, need attention, or failed."),
    ("07-import-manual-upload.png", "Import - Manual Upload", "This uploads the full MIS workbook or individual reports. The current run was built through the full workbook path."),
    ("08-sales-overview.png", "Sales - Overview", "This shows imported sales rows, platform totals, filters, and the sales records feeding the MIS."),
    ("09-sales-charts.png", "Sales - Charts", "This converts sales data into graphical platform/category visibility."),
    ("10-sales-platforms.png", "Sales - Platforms", "This breaks sales performance by platform."),
    ("11-sales-records.png", "Sales - Records", "This is the raw row-level audit trail for imported sales."),
    ("12-inventory-overview.png", "Inventory - Overview", "This connects sales quantities to inventory and stock movements."),
    ("13-inventory-stock.png", "Inventory - Stock", "This shows current stock position from inventory items and movements."),
    ("14-inventory-movements.png", "Inventory - Movements", "This is the stock ledger generated from sales sync and manual movements."),
    ("15-inventory-setup.png", "Inventory - Setup", "This is where inventory items and warehouse setup are maintained."),
    ("16-mis-preview.png", "MIS - Preview", "This is the main approval page for the calculated MIS result before final lock."),
    ("17-mis-charts.png", "MIS - Charts", "This provides graphical views of the MIS calculations, cost split, surplus, platform mix, and categories."),
    ("18-mis-profit-bridge.png", "MIS - Profit Bridge", "This explains movement from revenue to final net surplus/burn."),
    ("19-mis-platforms.png", "MIS - Platforms", "This shows platform-level MIS output."),
    ("20-mis-categories.png", "MIS - Categories", "This shows category profitability and gross-profit contribution."),
    ("21-mis-audit.png", "MIS - Audit", "This is the audit table behind the MIS output lines."),
    ("22-reports-overview.png", "Reports - Overview", "This groups the management reporting outputs."),
    ("23-reports-executive.png", "Reports - Executive", "This is the management summary for leadership review."),
    ("24-reports-pnl.png", "Reports - P&L", "This explains P&L-based cost mapping used by the MIS."),
    ("25-reports-loss-watch.png", "Reports - Loss Watch", "This highlights risk and loss areas."),
    ("26-validation.png", "Validation", "This explains notices, warnings, duplicates, unmapped items, and other audit checks."),
    ("27-adjustments.png", "Adjustments", "This allows manual monthly additions/deductions before recalculation."),
    ("28-masters.png", "Masters", "This maintains SKU mappings and COGS master data."),
]


def money(value) -> str:
    n = float(value or 0)
    sign = "-" if n < 0 else ""
    n = abs(n)
    return f"{sign}Rs. {n:,.2f}"


def num(value) -> str:
    if isinstance(value, str):
        try:
            value = float(value)
        except ValueError:
            return value
    return f"{value:,.2f}" if isinstance(value, float) else f"{value:,}"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, bold: bool = False, color: str | None = None) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    r = p.add_run(text)
    r.bold = bold
    if color:
        r.font.color.rgb = RGBColor.from_string(color)
    for paragraph in cell.paragraphs:
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = 1.08
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float] | None = None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Table Grid"
    for i, header in enumerate(headers):
        set_cell_text(t.rows[0].cells[i], header, True, "FFFFFF")
        set_cell_shading(t.rows[0].cells[i], "0F766E")
        if widths:
            t.rows[0].cells[i].width = Inches(widths[i])
    for row in rows:
        cells = t.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], str(value))
            if widths:
                cells[i].width = Inches(widths[i])
    doc.add_paragraph()
    return t


def add_callout(doc: Document, title: str, body: str, fill: str = "EAF8F6") -> None:
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = t.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    r = p.add_run(title)
    r.bold = True
    r.font.color.rgb = RGBColor(15, 118, 110)
    r.font.size = Pt(11)
    p2 = cell.add_paragraph(body)
    p2.paragraph_format.space_before = Pt(4)
    p2.paragraph_format.space_after = Pt(2)
    for paragraph in cell.paragraphs:
        paragraph.paragraph_format.left_indent = Pt(4)
        paragraph.paragraph_format.right_indent = Pt(4)
    doc.add_paragraph()


def add_image(doc: Document, image_path: Path, title: str, caption: str) -> None:
    doc.add_heading(title, level=3)
    p = doc.add_paragraph(caption)
    p.style = "Body Text"
    width = Inches(6.6)
    doc.add_picture(str(image_path), width=width)
    last = doc.paragraphs[-1]
    last.alignment = WD_ALIGN_PARAGRAPH.CENTER


def setup_styles(doc: Document) -> None:
    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(10)
    styles["Normal"].paragraph_format.space_after = Pt(6)
    for style_name, size, color in [
        ("Title", 24, "0B1220"),
        ("Heading 1", 17, "0F766E"),
        ("Heading 2", 14, "172033"),
        ("Heading 3", 11, "172033"),
    ]:
        st = styles[style_name]
        st.font.name = "Arial"
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = RGBColor.from_string(color)
        st.paragraph_format.space_before = Pt(10)
        st.paragraph_format.space_after = Pt(5)


def add_cover(doc: Document) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("MIS TOOL")
    r.bold = True
    r.font.size = Pt(12)
    r.font.color.rgb = RGBColor(15, 118, 110)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Architecture, Data Flow, Screens, and Calculation Guide")
    r.bold = True
    r.font.size = Pt(24)
    r.font.color.rgb = RGBColor(11, 18, 32)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Run ID 1 | February 2026 MIS")
    r.font.size = Pt(13)
    r.font.color.rgb = RGBColor(99, 112, 131)
    doc.add_paragraph()
    run = DATA["run"]
    table(doc, ["Item", "Value"], [
        ["Live URL", "http://localhost/miss/?run_id=1"],
        ["Run month", run["month"]],
        ["Run status", run["status"]],
        ["Last updated", run["updated_at"]],
        ["Imported sales rows", str(DATA["counts"]["import_rows"])],
        ["Final net surplus / burn", money(next(x["amount"] for x in DATA["overview"] if x["line_item"] == "Net surplus / burn"))],
    ], [2.0, 4.6])
    add_callout(doc, "How to read this document", "It starts from data entry/import, follows the backend tables and calculations, then ends at MIS preview, charts, reports, validation, and final approval output.")
    doc.add_page_break()


def add_architecture(doc: Document) -> None:
    doc.add_heading("1. End-to-End Tool Architecture", level=1)
    add_callout(doc, "Start to finish", "The tool starts with a monthly run. Data enters through manual workbook upload, portal automation, or API Connect. Everything is normalized into database tables, recalculated by the MIS calculator, reviewed in dashboards/reports, then finalized and exported.")
    doc.add_heading("1.1 Architecture Flow", level=2)
    table(doc, ["Stage", "Input", "Backend process", "Output screen"], [
        ["1. Create/select run", "Month/run ID", "monthly_runs row", "Dashboard"],
        ["2. Import data", "Excel, portal download, or API CSV/JSON", "Importer normalizes rows", "Import, Sales"],
        ["3. Map P&L and COGS", "Profit and loss / COGS Cal. sheets", "WorkbookMappingImporter stores cost mappings", "Masters, Reports P&L"],
        ["4. Calculate MIS", "Sales rows + P&L + COGS + adjustments", "MisCalculator refreshes summaries", "MIS Preview, MIS Charts"],
        ["5. Validate", "Imported and calculated data", "Validation rules create notices", "Validation Dashboard"],
        ["6. Review and finalize", "Management review", "Finalize/lock run", "Reports and export"],
    ], [1.25, 1.5, 2.25, 1.5])
    doc.add_heading("1.2 Main Code Modules", level=2)
    table(doc, ["Module", "What it does"], [
        ["src/App.php", "Routes, page rendering, POST actions, dashboard/report screens."],
        ["src/Database.php", "MySQL connection, migrations, table schema."],
        ["src/Services/Importer.php", "Reads uploaded reports/workbooks and inserts normalized sales rows."],
        ["src/Services/WorkbookMappingImporter.php", "Reads Profit and loss and COGS Cal. sheets."],
        ["src/Services/MisCalculator.php", "Builds platform summary, SKU summary, overview lines, validation, and final result."],
        ["src/Services/AutoImportService.php", "Creates and runs browser automation import jobs."],
        ["src/Services/ApiIntegrationService.php", "Stores API connections and imports CSV/JSON endpoints."],
        ["src/Services/InventoryService.php", "Syncs sales into warehouse/inventory movements."],
        ["assets/styles.css / assets/app.js", "Frontend styling, animations, counters, chart interactions, polling."],
    ], [2.35, 4.25])


def add_data_sources(doc: Document) -> None:
    doc.add_heading("2. Where Data Starts", level=1)
    doc.add_paragraph("The tool accepts data from three paths. The current run_id=1 was built from the full workbook upload path.")
    table(doc, ["Import path", "Routes", "What is captured", "Current run usage"], [
        ["Manual workbook upload", "/imports/manual, /imports/upload", "Sales sheets, Profit and loss, COGS Cal., SKU mappings, inventory items", "Used for run_id=1"],
        ["Browser portal automation", "/auto-import, /portal/connect, /imports/activity", "Downloaded marketplace reports after user login/OTP", "Available as fallback/automation"],
        ["API Connect", "/integrations, /integrations/save, /integrations/import", "CSV/JSON sales data from configured API endpoints", "Available for future direct API import"],
    ], [1.35, 1.6, 2.55, 1.1])
    doc.add_heading("2.1 Current Run Source Files", level=2)
    rows = [[s["source_type"], s["original_name"], str(s["rows_imported"]), s["import_mode"], s["uploaded_at"]] for s in DATA["sources"]]
    table(doc, ["Source", "File", "Rows", "Mode", "Uploaded"], rows[-5:], [1.2, 2.35, .7, .75, 1.25])
    doc.add_heading("2.2 Workbook Sheets and Columns", level=2)
    table(doc, ["Sheet", "Columns used", "Stored as", "Why it matters"], [
        ["Profit and loss", "A Account, B Amount, C P&L category, D Product category", "profit_loss_entries", "Column C maps costs into Marketing, Logistics, Seller Fee, COGS, Admin, etc."],
        ["COGS Cal.", "A item, B category, H multiplier, Y purchase price, AF packaging rate", "product_costs, sku_mappings, inventory_items", "Calculates raw material COGS, packaging, category mapping, and inventory master."],
        ["Sales/platform sheets", "Order, product, quantity, gross, tax, net, transaction fields", "import_rows", "Feeds platform totals, sales tables, category profitability, and MIS revenue."],
    ], [1.25, 2.0, 1.45, 1.9])


def add_calculations(doc: Document) -> None:
    doc.add_heading("3. How The MIS Calculates", level=1)
    add_callout(doc, "Important rule", "Screens do not recalculate directly from Excel. The importer saves normalized data first. MisCalculator then rebuilds summary tables. Pages read from those summary tables.")
    doc.add_heading("3.1 Formula Chain", level=2)
    table(doc, ["Step", "Formula / source"], [
        ["Sales including GST", "Sum of positive gross_amount from import_rows."],
        ["Returns", "Negative return/refund gross values from import_rows."],
        ["GST / tax", "Negative absolute sum of tax_amount."],
        ["Net sales after tax", "Sum of net_revenue from import_rows."],
        ["Platform costs", "P&L rows by category: Seller Fee, Logistics, Storage Charges, Transactions Charges, Packing, support, cold storage, labour."],
        ["Net proceeds", "Net sales after tax minus platform/operating costs."],
        ["Marketing spend", "P&L category Marketing."],
        ["After marketing", "Net proceeds minus marketing."],
        ["COGS - raw material", "quantity * multiplier * purchase_price from product_costs."],
        ["Packaging cost", "quantity * packaging_rate from product_costs."],
        ["Extra COGS from P&L", "P&L category COGS."],
        ["Gross margin after COGS", "After marketing minus raw material, packaging, and extra P&L COGS."],
        ["Net surplus / burn", "Gross margin after COGS minus agency/consultant and admin expenses."],
    ], [1.95, 4.65])
    doc.add_heading("3.2 Current MIS Output", level=2)
    rows = [[x["section"], x["line_item"], money(x["amount"]), f'{x["pct"]:.2f}%' if x["pct"] is not None else "", x["note"]] for x in DATA["overview"]]
    table(doc, ["Section", "Line item", "Amount", "%", "Reason"], rows, [1.0, 1.85, 1.05, .65, 2.05])


def add_results(doc: Document) -> None:
    doc.add_heading("4. Current Run Results", level=1)
    doc.add_heading("4.1 Platform Data", level=2)
    table(doc, ["Platform", "Rows", "Qty", "Gross", "Tax", "Net"], [
        [p["platform"], str(p["rows_count"]), num(float(p["qty"])), money(p["gross"]), money(p["tax"]), money(p["net"])] for p in DATA["platforms"]
    ], [1.55, .65, .65, 1.15, 1.0, 1.15])
    doc.add_heading("4.2 Top Category Profitability", level=2)
    table(doc, ["Category", "Qty", "Revenue", "COGS", "Packaging", "Gross profit"], [
        [c["category"], num(float(c["qty"])), money(c["revenue"]), money(c["cogs"]), money(c["packaging"]), money(c["gross_profit"])] for c in DATA["categories"][:8]
    ], [2.0, .55, 1.05, .9, .9, 1.05])
    doc.add_heading("4.3 P&L Mapping Categories", level=2)
    table(doc, ["P&L category", "Rows", "Amount", "Used for"], [
        [p["pnl_category"], str(p["rows_count"]), money(p["amount"]), pnl_use(p["pnl_category"])] for p in DATA["pnl"]
    ], [1.8, .55, 1.1, 3.15])


def pnl_use(cat: str) -> str:
    m = {
        "Marketing": "Marketing spend",
        "Logistics": "Fulfilment and logistics",
        "Storage Charges": "Storage charges",
        "Seller Fee": "Selling fee / commission",
        "COGS": "Extra COGS from P&L",
        "Packing": "Packing cost line",
        "Transactions Charges": "Payment transaction charges",
        "Other support services": "Other support services",
        "Cold Storage Charges": "Cold storage charges",
        "Professional fees": "General/professional expenses",
        "G & A expenses": "General/admin expenses",
        "Code incentive": "Agency / consultant fees",
        "Snell Business collective LLP": "Agency / consultant fees",
        "Muskaan jain": "Agency / consultant fees",
        "Misc. Expense": "General/professional expenses",
    }
    return m.get(cat, "Review/mapped cost bucket")


def add_database_and_routes(doc: Document) -> None:
    doc.add_heading("5. Backend Tables and Outputs", level=1)
    table(doc, ["Table", "Purpose"], [
        ["monthly_runs", "One MIS month/run, status, lock/finalization state."],
        ["source_files", "Uploaded/captured source files and import row counts."],
        ["import_rows", "Normalized sales/return rows from all sources."],
        ["profit_loss_entries", "P&L rows mapped from workbook column C."],
        ["product_costs", "COGS multiplier, purchase price, packaging rates."],
        ["sku_mappings", "Product to COGS SKU/MIS category mapping."],
        ["inventory_items", "Inventory master items from COGS sheet."],
        ["inventory_movements", "Stock ledger generated from sales sync and movements."],
        ["mis_platform_summary", "Platform-level calculated summary."],
        ["mis_sku_summary", "Category/platform revenue, COGS, packaging, gross profit."],
        ["mis_overview_lines", "Final MIS lines shown in dashboards and reports."],
        ["validation_issues", "Notices, warnings, errors for review."],
        ["monthly_adjustments", "Manual additions/deductions before recalculation."],
        ["auto_import_jobs", "Browser automation job records and logs."],
        ["api_integrations", "Saved API connector configuration."],
    ], [2.0, 4.6])
    doc.add_heading("5.1 Route Map", level=2)
    table(doc, ["Area", "Routes"], [
        ["Dashboard", "/, /dashboard/close, /dashboard/trends"],
        ["Import", "/auto-import, /integrations, /imports/activity, /imports/manual"],
        ["Sales", "/sales, /sales/charts, /sales/platforms, /sales/records"],
        ["Inventory", "/inventory, /inventory/stock, /inventory/movements, /inventory/setup"],
        ["MIS", "/mis/preview, /mis/charts, /mis/profit-bridge, /mis/platforms, /mis/categories, /mis/audit"],
        ["Reports", "/reports, /reports/executive, /reports/pnl, /reports/loss-watch"],
        ["Control", "/validation, /adjustments, /masters"],
    ], [1.4, 5.2])


def add_screens(doc: Document) -> None:
    doc.add_heading("6. Screen-by-Screen Walkthrough", level=1)
    add_callout(doc, "How screenshots are ordered", "The screenshots follow the real workflow: start at dashboard, import data, review sales/inventory, inspect MIS, review reports and validation, then finalize or adjust.")
    for idx, (filename, title, caption) in enumerate(SCREENSHOTS, 1):
        add_image(doc, SHOT_DIR / filename, f"6.{idx} {title}", caption)
        if idx in {3, 7, 11, 15, 21, 25}:
            doc.add_page_break()


def add_end(doc: Document) -> None:
    doc.add_heading("7. Where The Tool Ends", level=1)
    doc.add_paragraph("The workflow ends when MIS Preview and reports are reviewed, validation notices are accepted or fixed, adjustments are applied if needed, and the run is finalized/locked.")
    table(doc, ["Final stage", "What user checks", "Output"], [
        ["MIS Preview", "Revenue, costs, category margin, platform totals, final surplus/burn.", "Approval-ready MIS."],
        ["Validation", "Duplicate files, missing mappings, unusual rows, notices.", "Confidence before finalization."],
        ["Adjustments", "Manual additions/deductions if month-end accounting requires them.", "Recalculated MIS."],
        ["Reports", "Executive summary, P&L report, loss watch.", "Management presentation view."],
        ["Export/finalize", "Lock once approved.", "Stable monthly MIS output."],
    ], [1.35, 3.15, 2.1])
    add_callout(doc, "Current final result for run_id=1", "Net sales after tax is Rs. 19,52,780.04 and final net surplus / burn is Rs. 3,58,580.28, equal to 18.36% of net sales.", "FFF7DF")


def main() -> None:
    doc = Document()
    setup_styles(doc)
    sec = doc.sections[0]
    sec.top_margin = Inches(0.7)
    sec.bottom_margin = Inches(0.7)
    sec.left_margin = Inches(0.75)
    sec.right_margin = Inches(0.75)
    add_cover(doc)
    add_architecture(doc)
    add_data_sources(doc)
    add_calculations(doc)
    add_results(doc)
    add_database_and_routes(doc)
    add_screens(doc)
    add_end(doc)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
